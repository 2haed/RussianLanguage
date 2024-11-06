import os
import tempfile
from io import BytesIO

import aspose.words as aw
from docx import Document


def convert_doc_to_docx_aspose(input_file, output_file):
    doc = aw.Document(input_file)
    doc.save(output_file)


def remove_evaluation_text(docx_file):
    doc = Document(docx_file)
    evaluation_text = "Created with an evaluation copy of Aspose.Words. To remove all limitations, you can use Free Temporary License https://products.aspose.com/words/temporary-license/"

    for paragraph in doc.paragraphs:
        if evaluation_text in paragraph.text:
            p = paragraph._element
            p.getparent().remove(p)
    doc.save(docx_file)


def process_file(file_content, file_extension):
    text_content = ""
    with tempfile.TemporaryDirectory() as temp_dir:
        if file_extension == '.txt':
            text_content = file_content.getvalue().decode('utf-8')

        elif file_extension == ".docx":
            temp_file_path = os.path.join(temp_dir, 'temp.docx')
            with open(temp_file_path, 'wb') as f:
                f.write(file_content.getvalue())

            # Чтение docx файла
            doc = Document(temp_file_path)
            text_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        elif file_extension == ".doc":
            temp_doc_path = os.path.join(temp_dir, 'temp.doc')
            temp_docx_path = os.path.join(temp_dir, 'temp.docx')

            with open(temp_doc_path, 'wb') as f:
                f.write(file_content.getvalue())

            # Конвертация через Aspose.Words
            convert_doc_to_docx_aspose(temp_doc_path, temp_docx_path)

            # Удаление текста об оценочной копии
            remove_evaluation_text(temp_docx_path)

            # Чтение сконвертированного docx файла
            doc = Document(temp_docx_path)
            text_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        # Временные файлы автоматически удалятся при выходе из блока with

    return text_content

file_path = os.path.abspath('word.doc')

# Получаем расширение файла
file_extension = os.path.splitext(file_path)[1].lower()

# Открываем файл и загружаем его содержимое
with open(file_path, 'rb') as f:
    file_content = BytesIO(f.read())

# Обрабатываем файл с помощью функции process_file_with_aspose
try:
    text_content = process_file(file_content, file_extension)
    print("Содержимое файла:")
    print(text_content)
except Exception as e:
    print(f"Произошла ошибка при обработке файла: {str(e)}")